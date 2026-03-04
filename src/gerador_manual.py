"""
Gerador Automático de Manuais Word
Baseado no padrão corporativo definido
Versão: 2.0.0 - Padrão Profissional
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def criar_manual(json_path: str, output_path: str):
    """
    Gera manual .docx a partir de JSON estruturado
    Implementa padrão profissional com breadcrumbs, formatação avançada e rodapé detalhado
    
    Args:
        json_path: Caminho do JSON com conteúdo
        output_path: Caminho de saída do .docx
    """
    
    # Carregar dados
    with open(json_path, 'r', encoding='utf-8') as f:
        dados = json.load(f)
    
    # Criar documento
    doc = Document()
    
    # ===== CAPA PROFISSIONAL =====
    criar_capa_profissional(doc, dados['metadata'], json_path)
    doc.add_page_break()
    
    # ===== SUMÁRIO =====
    criar_sumario(doc)
    doc.add_page_break()
    
    # ===== OBJETIVO =====
    doc.add_heading('1. Objetivo', level=1)
    doc.add_paragraph(dados['objetivo'])
    doc.add_paragraph()
    
    # ===== PRÉ-REQUISITO =====
    doc.add_heading('2. Pré-requisito', level=1)
    doc.add_paragraph(dados['pre_requisito'])
    doc.add_paragraph()
    
    # ===== FUNCIONALIDADES =====
    doc.add_heading('3. Funcionalidade', level=1)
    
    for idx, func in enumerate(dados['funcionalidades'], start=1):
        # Breadcrumb de navegação
        breadcrumb = doc.add_paragraph()
        breadcrumb.paragraph_format.left_indent = Inches(0.25)
        bc_run = breadcrumb.add_run(f'{dados["metadata"]["modulo"]} >> {func["titulo"]}')
        bc_run.font.size = Pt(10)
        bc_run.font.color.rgb = RGBColor(100, 100, 150)
        bc_run.italic = True
        
        # Título da funcionalidade com numeração hierárquica
        doc.add_heading(f'3.{idx} {func["titulo"]}', level=2)
        
        # Prints (se houver)
        json_dir = Path(json_path).parent
        if func.get('prints'):
            for print_idx, print_path in enumerate(func['prints'], 1):
                full_path = json_dir / print_path
                try:
                    doc.add_picture(str(full_path), width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Legenda da imagem
                    legenda = doc.add_paragraph()
                    legenda_run = legenda.add_run(f'Figura 3.{idx}.{print_idx}: {func["titulo"]}')
                    legenda_run.font.size = Pt(9)
                    legenda_run.font.italic = True
                    legenda_run.font.color.rgb = RGBColor(100, 100, 100)
                    legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()  # Espaçamento
                except Exception as e:
                    print(f"[AVISO] Erro ao inserir imagem {print_path}: {e}")
                    doc.add_paragraph(f'[Imagem não encontrada: {print_path}]')
        
        # Descrição
        doc.add_paragraph(func['descricao'])
        
        # Observações formatadas
        if func.get('observacoes'):
            obs_heading = doc.add_paragraph()
            obs_heading.paragraph_format.left_indent = Inches(0.25)
            obs_run = obs_heading.add_run('Observações:')
            obs_run.bold = True
            obs_run.font.size = Pt(11)
            
            for obs_idx, obs in enumerate(func['observacoes'], 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p_run = p.add_run(f'• {obs}')
                p_run.font.size = Pt(10)
        
        # Ícones com legendas (se fornecidos)
        if func.get('icones'):
            icons_heading = doc.add_paragraph()
            icons_heading.paragraph_format.left_indent = Inches(0.25)
            icons_run = icons_heading.add_run('Ícones Utilizados:')
            icons_run.bold = True
            icons_run.font.size = Pt(11)
            
            for icone in func['icones']:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p_run = p.add_run(f'• {icone["nome"]}: {icone["descricao"]}')
                p_run.font.size = Pt(10)
        
        doc.add_paragraph()
    
    # ===== RODAPÉ PROFISSIONAL =====
    aplicar_rodape_profissional(doc, dados['metadata'])
    
    # Garantir diretório de saída
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    # Salvar
    doc.save(output_path)
    print(f'[OK] Manual gerado com sucesso: {output_path}')


def criar_capa_profissional(doc, metadata, json_path):
    """Cria capa padronizada com design profissional"""
    
    # Logo (se existir)
    if metadata.get('logo_path'):
        json_dir = Path(json_path).parent
        logo_full_path = json_dir / metadata['logo_path']
        try:
            doc.add_picture(str(logo_full_path), width=Inches(1.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"[AVISO] Logo não encontrada: {e}")
    
    # Espaçamento
    for _ in range(4):
        doc.add_paragraph()
    
    # Título principal
    titulo = doc.add_heading(metadata['nome_manual'], level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.size = Pt(28)
    titulo.runs[0].font.bold = True
    titulo.runs[0].font.color.rgb = RGBColor(30, 30, 30)
    
    # Subtítulo com módulo
    subtitulo = doc.add_paragraph(metadata['modulo'])
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].font.size = Pt(14)
    subtitulo.runs[0].font.bold = True
    subtitulo.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    
    # Sistema (se fornecido)
    if metadata.get('sistema'):
        sistema_p = doc.add_paragraph(metadata['sistema'])
        sistema_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sistema_p.runs[0].font.size = Pt(12)
        sistema_p.runs[0].font.color.rgb = RGBColor(120, 120, 120)
    
    # Espaçamento
    for _ in range(6):
        doc.add_paragraph()
    
    # Seção de metadados
    metadata_section = doc.add_paragraph()
    metadata_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tabela de informações de rodapé da capa
    info_text = (
        f"Elaborado: {metadata['elaborado']}\n"
        f"Revisado: {metadata['revisado']}\n"
        f"Classificação: {metadata['classificacao'].upper()}"
    )
    
    for line in info_text.split('\n'):
        info_p = doc.add_paragraph(line)
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_p.runs[0].font.size = Pt(10)
        info_p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    # Linha decorativa
    linha = doc.add_paragraph()
    linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    linha_run = linha.add_run('─' * 50)
    linha_run.font.color.rgb = RGBColor(150, 150, 150)


def criar_sumario(doc):
    """Cria página de sumário com campo TOC do Word"""
    titulo = doc.add_heading('Sumário', level=1)
    titulo.runs[0].font.size = Pt(18)
    
    # Criar campo TOC (Table of Contents) do Word
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Início do campo
    fldChar1 = run._element.makeelement(qn('w:fldChar'))
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    # Instrução do campo TOC
    instrText = run._element.makeelement(qn('w:instrText'))
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    run._element.append(instrText)

    # Separador SEPARATE
    fldChar2 = run._element.makeelement(qn('w:fldChar'))
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)

    # Texto placeholder
    run2 = paragraph.add_run('Clique com botão direito e selecione "Atualizar campo" para gerar o sumário...')
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(150, 150, 150)
    run2.font.size = Pt(10)

    # Fim do campo
    run3 = paragraph.add_run()
    fldChar3 = run3._element.makeelement(qn('w:fldChar'))
    fldChar3.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar3)


def aplicar_rodape_profissional(doc, metadata):
    """Aplica rodapé padronizado e profissional em todas as páginas"""
    section = doc.sections[0]
    footer = section.footer
    
    # Limpar parágrafo existente
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()
    
    # Linha separadora do rodapé
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '999999')
    pBdr.append(top)
    pPr.append(pBdr)
    
    # Limpar conteúdo anterior
    for i in range(len(p.runs)):
        r = p.runs[0]._element
        r.getparent().remove(r)
    
    # Construir rodapé com informações detalhadas
    # Parte 1: Sistema e Módulo
    sistema_texto = metadata.get('sistema', 'Sistema')
    
    # Usar runs para montar o footer
    footer_content = (
        f"Elaborado: {metadata['elaborado']} • "
        f"Revisado: {metadata['revisado']} • "
        f"Classificação: {metadata['classificacao'].upper()} • "
        f"Página "
    )
    
    run_text = p.add_run(footer_content)
    run_text.font.size = Pt(8)
    run_text.font.color.rgb = RGBColor(100, 100, 100)
    
    # Adicionar número de página atual
    run = p.add_run()
    fldChar1 = run._element.makeelement(qn('w:fldChar'))
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    instrText = run._element.makeelement(qn('w:instrText'))
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._element.append(instrText)
    
    fldChar2 = run._element.makeelement(qn('w:fldChar'))
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)
    
    run_slash = p.add_run(' de ')
    run_slash.font.size = Pt(8)
    run_slash.font.color.rgb = RGBColor(100, 100, 100)
    
    # Adicionar número total de páginas
    run = p.add_run()
    fldChar1 = run._element.makeelement(qn('w:fldChar'))
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    instrText = run._element.makeelement(qn('w:instrText'))
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    run._element.append(instrText)
    
    fldChar2 = run._element.makeelement(qn('w:fldChar'))
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)
    
    # Centralizar e formatar
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    """Interface CLI"""
    if len(sys.argv) != 3:
        print("Uso: python gerador_manual.py <input.json> <output.docx>")
        print("\nExemplo:")
        print("  python gerador_manual.py exemplos/input/manual_input.json exemplos/output/Manual.docx")
        sys.exit(1)
    
    json_path = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        criar_manual(json_path, output_path)
    except FileNotFoundError as e:
        print(f"[ERRO] {e}")
        sys.exit(1)
    except json.JSONDecodeError:
        print(f"[ERRO] JSON invalido em {json_path}")
        sys.exit(1)
    except Exception as e:
        print(f"[ERRO] Ao gerar manual: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
